"""
Export word search puzzles to DOCX format.

This module provides functionality to save word search puzzles
as Microsoft Word documents (.docx files).
"""

import os
from docx import Document


def save_wordsearch_to_docx(file_path, title, grid, words, solution=None):
    """
    Saves the wordsearch result to a DOCX file.

    Args:
        file_path (str): The full path where the DOCX file will be saved.
        title (str): The title of the word search puzzle.
        grid (list): The puzzle grid as a 2D list of letters.
        words (list): The list of words in the puzzle.
    """
    doc = Document()
    doc.add_heading(title, 0)

    # Add the grid
    doc.add_heading("Puzzle Grid", level=1)
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    for i, row in enumerate(grid):
        for j, letter in enumerate(row):
            table.cell(i, j).text = letter

    # Add the word list
    doc.add_heading("Word List", level=1)
    for word in words:
        doc.add_paragraph(word, style="List Bullet")

    # [TO DO]: add solution page, with a dedicated parameter.
    if solution:
        pass

    # Ensure the output directory exists and save the document
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)


if __name__ == "__main__":
    # Example usage
    wordsearch_result = {
        "grid": [
            ["A", "B", "C", "D"],
            ["E", "F", "G", "H"],
            ["I", "J", "K", "L"],
            ["M", "N", "O", "P"],
        ],
        "words": ["ABCD", "EFGH", "IJKL", "MNOP"],
    }
    output_path = os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "..", "out", "wordsearch.docx")
    )
    save_wordsearch_to_docx(
        output_path,
        "Word Search Puzzle Title",
        wordsearch_result["grid"],
        wordsearch_result["words"],
    )
